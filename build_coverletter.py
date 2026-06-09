from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Date
p = doc.add_paragraph()
run = p.add_run(f"Date: {date.today().strftime('%B %d, %Y')}")
run.font.size = Pt(10)

# Recipient
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
lines = [
    "Hiring Manager",
    "TestMu AI (formerly LambdaTest)",
    "1 Sutter Street, Suite 500",
    "San Francisco, California 94104",
]
for i, line in enumerate(lines):
    if i > 0:
        p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.size = Pt(10)

# Subject
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("Re: Application for Solutions Engineer — Noida")
run.bold = True
run.font.size = Pt(11)

# Salutation
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
run = p.add_run("Dear Hiring Manager,")
run.font.size = Pt(10)

# Body
body_paragraphs = [
    (
        "I am writing to express my interest in the Solutions Engineer role at TestMu AI. "
        "With over 3 years of experience in SaaS product engineering at RxLogix — a regulated "
        "pharmacovigilance platform serving FDA, EMA, Merck, and J&J — I have developed a strong "
        "foundation in technical problem-solving, cross-functional collaboration, and delivering "
        "customer-facing engineering outcomes."
    ),
    (
        "The role's focus on Customer Trial Management resonates directly with my experience. "
        "At RxLogix, I managed end-to-end client-specific releases, service packs, and hotfixes "
        "for enterprise pharma clients, working closely with the Consulting team to translate "
        "customer requirements into technical deliverables. I monitored JIRA tickets from client-facing "
        "consultants, performed RCA on production issues, and ensured timely delivery — skills that "
        "directly map to owning and driving successful customer trials."
    ),
    (
        "On the technical side, I have hands-on experience with Java, Python, REST APIs, "
        "PostgreSQL, and Oracle DB. I have debugged complex production issues across the full "
        "stack — from JVM heap dumps to SQL query optimization — and have a strong understanding "
        "of SDLC in Agile/Scrum environments. I have also adopted AI-assisted development tools "
        "(GitHub Copilot, Claude Code) to improve velocity, which aligns with TestMu AI's mission "
        "of building AI-native testing platforms."
    ),
    (
        "I am based in Noida/Delhi and am excited about the opportunity to bridge technical depth "
        "with customer success at an AI-first company like TestMu AI. I would welcome the chance "
        "to discuss how my background can contribute to your team."
    ),
]

for text in body_paragraphs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)

# Sign-off
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Best regards,")
run.font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
run = p.add_run("Rishabh Rajpurohit")
run.bold = True
run.font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
run = p.add_run("+91 8929566279 | rajpurohitrishabh1@gmail.com")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

p = doc.add_paragraph()
run = p.add_run("github.com/rishxbhhhh | linkedin.com/in/rishabhrajpurohit1")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

output = "C:/Users/rishabh/Documents/Resume/CoverLetter_TestMuAI_SolutionsEngineer.docx"
doc.save(output)
print(f"Saved: {output}")
