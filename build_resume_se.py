from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# ── Style helpers ──
NAME_COLOR = RGBColor(0x1A, 0x5F, 0x7A)
HEADING_COLOR = RGBColor(0x1A, 0x5F, 0x7A)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A5F7A')
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = HEADING_COLOR
    add_horizontal_rule(doc)

def job_header(doc, title, company, dates):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    left.paragraphs[0].add_run(title).bold = True
    left.paragraphs[0].runs[0].font.size = Pt(9.5)
    left.paragraphs[0].add_run(f"  —  {company}")
    left.paragraphs[0].runs[1].font.size = Pt(9.5)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraphs[0].add_run(dates)
    right.paragraphs[0].runs[0].font.size = Pt(9.5)
    right.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x5F, 0x63, 0x68)
    # Remove table borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def role_header(doc, title, dates):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    run = left.paragraphs[0].add_run(title)
    run.italic = True
    run.font.size = Pt(9)
    left.paragraphs[0].paragraph_format.left_indent = Inches(0.15)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = right.paragraphs[0].add_run(dates)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)
    right.paragraphs[0].runs[0].font.size = Pt(9)
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for bn in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{bn}')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            tcPr.append(tcBorders)

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    for run in p.runs:
        run.font.size = Pt(9)

# ═══════════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Rishabh Rajpurohit")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = NAME_COLOR

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Solutions Engineer | Technical Customer Success | SaaS Trials")
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
run = p.add_run("+91 8929566279  |  New Delhi, India  |  rajpurohitrishabh1@gmail.com")
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run("linkedin.com/in/rishabhrajpurohit1  |  github.com/rishxbhhhh")
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

# ═══════════════════════════════════════════════
# PROFESSIONAL SUMMARY
# ═══════════════════════════════════════════════
section_heading(doc, "Professional Summary")

summary_text = (
    "Results-driven engineer with 3+ years of experience in SaaS product engineering, "
    "customer-facing technical delivery, and cross-functional collaboration at RxLogix, "
    "a GxP-regulated pharmacovigilance platform serving FDA, EMA, Merck, and J&J. "
    "Proven track record of managing end-to-end client-specific releases, translating "
    "customer requirements into technical solutions, and driving trial-to-production "
    "success. Strong technical foundation in Java, Python, REST APIs, databases, and "
    "Agile SDLC. Adept at bridging the gap between engineering and client stakeholders — "
    "capturing feedback, troubleshooting complex issues, and ensuring seamless product adoption."
)
p = doc.add_paragraph(summary_text)
p.paragraph_format.space_after = Pt(4)
for run in p.runs:
    run.font.size = Pt(9)

# ═══════════════════════════════════════════════
# EXPERIENCE
# ═══════════════════════════════════════════════
section_heading(doc, "Experience")

# Company header
job_header(doc, "RxLogix Corporation", "Pharmacovigilance SaaS", "Mar 2023 – Present · 3 yrs 3 mos")

role_header(doc, "Software Engineer · Promoted from ASE II  |  Customer-Facing Engineering", "Dec 2025 – Present")

bullet(doc, "Managed end-to-end client-specific releases, service packs, and hotfixes for enterprise pharma clients (FDA, EMA, Merck, Bayer, J&J), ensuring on-time delivery and client satisfaction")
bullet(doc, "Collaborated cross-functionally with AI, UI, and QA teams to deliver AI Early Detection widgets — translated product requirements into backend APIs and coordinated with internal teams to ensure customer trial success")
bullet(doc, "Delivered production bug fixes and feature enhancements tracked via JIRA (KANBAN + Gantt), maintaining clear communication with stakeholders on timelines and scope")
bullet(doc, "Provided technical troubleshooting and RCA for client-reported issues across Groovy/Grails, Oracle DB, PostgreSQL, and REST API layers")
bullet(doc, "Adopted AI-assisted development tools (GitHub Copilot, Claude Code) to accelerate feature velocity and improve code review quality")

role_header(doc, "Associate Software Engineer II · Promoted from ASE I  |  Client Support & Delivery", "Jan 2024 – Nov 2025")

bullet(doc, "Monitored JIRA tickets from the Consulting team, performed root cause analysis, and resolved backend issues — acting as the technical bridge between client-facing consultants and engineering")
bullet(doc, "Delivered hotfixes and patch releases for client-specific deployments in a GxP-compliant environment with strict audit and validation requirements")
bullet(doc, "Migrated application database from Oracle to PostgreSQL while maintaining zero-downtime for ongoing client trials")
bullet(doc, "Used SQL Developer, PG Admin, and Postman for client-data troubleshooting, schema root-cause analysis, and API verification")

role_header(doc, "Associate Software Engineer I · First Role", "Mar 2023 – Jan 2024")

bullet(doc, "Supported the product engineering team with application debugging, bug triage, and JVM heap analysis using Eclipse MAT")
bullet(doc, "Learned Groovy, Grails, and the full SDLC lifecycle under Agile/Scrum methodology in a regulated environment")

# ═══════════════════════════════════════════════
# SKILLS
# ═══════════════════════════════════════════════
section_heading(doc, "Skills")

skills_data = [
    ("Languages", "Groovy, Java, Python, SQL"),
    ("Frameworks", "Grails, Spring Boot, GORM/Hibernate"),
    ("Databases", "Oracle DB, PostgreSQL, SQLite"),
    ("Dev & Tools", "Git, JIRA, KANBAN, IntelliJ, Postman, Eclipse MAT"),
    ("SDLC & Practices", "Agile/Scrum, RCA, Hotfix Delivery, GxP Compliance"),
    ("Customer Skills", "Trial Management, Technical Troubleshooting, Cross-Functional Collaboration"),
    ("AI Tooling", "GitHub Copilot, Claude Code"),
]

table = doc.add_table(rows=len(skills_data), cols=2)
table.autofit = True

for i, (category, value) in enumerate(skills_data):
    left = table.cell(i, 0)
    right = table.cell(i, 1)
    left.paragraphs[0].add_run(category).bold = True
    left.paragraphs[0].runs[0].font.size = Pt(9)
    right.paragraphs[0].add_run(value)
    right.paragraphs[0].runs[0].font.size = Pt(9)
    # No borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for bn in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{bn}')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            tcPr.append(tcBorders)

# ═══════════════════════════════════════════════
# PROJECTS
# ═══════════════════════════════════════════════
section_heading(doc, "Projects")

job_header(doc, "CipherChat", "Secure REST Chat Backend", "")
bullet(doc, "Built secure chat backend with Spring Boot, AES-at-rest encryption, and RSA hybrid key sharing")
bullet(doc, "Implemented RESTful APIs, Liquibase schema migrations, and CORS for cross-origin access")
bullet(doc, "GitHub: github.com/rishxbhhhh/cipherchat")

job_header(doc, "CheckStock", "Real-Time Telegram Bot", "")
bullet(doc, "Developed a Python-based Telegram bot for commodity stock monitoring with multi-user support")
bullet(doc, "Features: configurable polling, bot-controlled pause/resume, real-time async notifications")
bullet(doc, "GitHub: github.com/rishxbhhhh/check_stock")

# ═══════════════════════════════════════════════
# EDUCATION
# ═══════════════════════════════════════════════
section_heading(doc, "Education")

job_header(doc, "B.Tech, Computer Science & Engineering", "Jaypee Institute of Information Technology, Noida", "Jul 2019 – Mar 2023")
p = doc.add_paragraph("CGPA: 7.1 / 10 | Relevant: Data Structures, Algorithms, DBMS, OOP, Computer Networks")
for run in p.runs:
    run.font.size = Pt(9)

# Save
output = "C:/Users/rishabh/Documents/Resume/Rishabh_Rajpurohit_TestMuAI_SolutionsEngineer.docx"
doc.save(output)
print(f"Saved: {output}")
